import os

from datetime import datetime

import pandas as pd

from models.main_model import MainModel

class MainController:

    def __init__(self):

        self.model = MainModel()

    def get_transaction_list(self):

        return self.model.get_all_transactions()

    def clean_amount(self, amount):

        return amount.replace(",", "").strip()

    def validate_data(self, date, description, amount, category):

        if date.strip() == "":

            return False, "Vui lòng chọn ngày!"

        if description.strip() == "":

            return False, "Vui lòng nhập nội dung!"

        if amount.strip() == "":

            return False, "Vui lòng nhập số tiền!"

        if category.strip() == "":

            return False, "Vui lòng chọn danh mục!"

        amount_clean = self.clean_amount(amount)

        try:

            amount_float = float(amount_clean)

            if amount_float <= 0:

                return False, "Số tiền phải lớn hơn 0!"

        except ValueError:

            return False, "Số tiền phải là kiểu số!"

        return True, "Dữ liệu hợp lệ!"

    def add_transaction(self, transaction_type, date, description, amount, category):

        valid, message = self.validate_data(date, description, amount, category)

        if not valid:

            return False, message

        amount_clean = self.clean_amount(amount)

        self.model.add_transaction(transaction_type, date, description, amount_clean, category)

        return True, "Thêm giao dịch thành công!"

    def update_transaction(self, transaction_id, transaction_type, date, description, amount, category):

        valid, message = self.validate_data(date, description, amount, category)

        if not valid:

            return False, message

        amount_clean = self.clean_amount(amount)

        self.model.update_transaction(transaction_id, transaction_type, date, description, amount_clean, category)

        return True, "Sửa giao dịch thành công!"

    def delete_transaction(self, transaction_id):

        self.model.delete_transaction(transaction_id)

        return True, "Xóa giao dịch thành công!"

    def import_transactions_from_csv(self, file_path, replace_data=False):

        imported_count = self.model.import_csv_file(file_path, replace_data)

        return True, f"Nhập CSV thành công! Đã thêm {imported_count} giao dịch."

    def export_transactions_to_csv(self, file_path, transactions=None):

        self.model.export_csv_file(file_path, transactions)

        return True, "Xuất CSV thành công!"

    def get_total_income(self):

        return self.model.get_total_income()

    def get_total_expense(self):

        return self.model.get_total_expense()

    def get_balance(self):

        return self.model.get_balance()

    def search_transactions(

        self,

        keyword="",

        transaction_type="Tất cả",

        category="Tất cả",

        date_from=None,

        date_to=None,

        min_amount=None,

        max_amount=None

    ):

        return self.model.search_transactions(

            keyword,

            transaction_type,

            category,

            date_from,

            date_to,

            min_amount,

            max_amount

        )

    def get_expenses_by_category(self, transactions=None):

        return self.model.get_expenses_by_category(transactions)

    def get_category_expense_percentages(self, transactions=None):

        return self.model.get_expense_percentage_by_category(transactions)

    def get_monthly_expense_summary(self, transactions=None):

        return self.model.get_expenses_by_month(transactions)

    def generate_report(self, transactions=None):

        if transactions is None:

            transactions = self.model.get_all_transactions()

        total_income = sum(self.model.to_float(t["amount"]) for t in transactions if t.get("type") == "Thu nhập")

        total_expense = sum(self.model.to_float(t["amount"]) for t in transactions if t.get("type") == "Chi tiêu")

        balance = total_income - total_expense

        by_category = self.model.get_expense_percentage_by_category(transactions)

        by_month = self.model.get_expenses_by_month(transactions)

        lines = [

            "BÁO CÁO THU CHI",

            "===========================",

            f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}",

            "",

            f"Tổng thu nhập: {int(total_income):,} đ",

            f"Tổng chi tiêu: {int(total_expense):,} đ",

            f"Số dư còn lại: {int(balance):,} đ",

            "",

            "Chi tiêu theo danh mục:",

        ]

        if by_category:

            for category, (amount, percent) in sorted(by_category.items(), key=lambda x: x[1][0], reverse=True):

                lines.append(f" - {category}: {int(amount):,} đ ({percent}%)")

        else:

            lines.append(" - Chưa có giao dịch chi tiêu")

        lines.append("")

        lines.append("Chi tiêu theo tháng:")

        if by_month:

            for month, amount in sorted(by_month.items(), key=lambda x: datetime.strptime(x[0], "%m/%Y")):

                lines.append(f" - {month}: {int(amount):,} đ")

        else:

            lines.append(" - Chưa có giao dịch chi tiêu")

        lines.append("")

        lines.append("Danh sách giao dịch:")

        if transactions:

            for transaction in sorted(transactions, key=lambda t: self.model.parse_date(t.get("date", "")) or datetime.min):

                lines.append(

                    f"{transaction.get('date', '')} | {transaction.get('type', '')} | {transaction.get('category', 'Khác')} | "

                    f"{transaction.get('description', '')} | {int(self.model.to_float(transaction.get('amount', '0'))):,} đ"

                )

        else:

            lines.append(" - Không có giao dịch nào trong bộ lọc hiện tại")

        return "\n".join(lines)

    def export_report_to_word(self, file_path, transactions=None):

        from docx import Document

        from docx.shared import Pt

        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if transactions is None:

            transactions = self.model.get_all_transactions()

        total_income = sum(self.model.to_float(t["amount"]) for t in transactions if t.get("type") == "Thu nhập")

        total_expense = sum(self.model.to_float(t["amount"]) for t in transactions if t.get("type") == "Chi tiêu")

        balance = total_income - total_expense

        category_percentages = self.model.get_expense_percentage_by_category(transactions)

        monthly_expenses = self.model.get_expenses_by_month(transactions)

        document = Document()

        title = document.add_heading("BÁO CÁO THU CHI CÁ NHÂN", level=1)

        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

        document.add_heading("1. Tổng quan tài chính", level=2)

        document.add_paragraph(f"Tổng thu nhập: {int(total_income):,} đ")

        document.add_paragraph(f"Tổng chi tiêu: {int(total_expense):,} đ")

        document.add_paragraph(f"Số dư còn lại: {int(balance):,} đ")

        document.add_heading("2. Chi tiêu theo danh mục", level=2)

        if category_percentages:

            for category, (amount, percent) in sorted(category_percentages.items(), key=lambda x: x[1][0], reverse=True):

                document.add_paragraph(f"{category}: {int(amount):,} đ ({percent}%)", style="List Bullet")

        else:

            document.add_paragraph("Chưa có giao dịch chi tiêu", style="List Bullet")

        document.add_heading("3. Chi tiêu theo tháng", level=2)

        if monthly_expenses:

            for month, amount in sorted(monthly_expenses.items(), key=lambda x: datetime.strptime(x[0], "%m/%Y")):

                document.add_paragraph(f"{month}: {int(amount):,} đ", style="List Bullet")

        else:

            document.add_paragraph("Chưa có giao dịch chi tiêu", style="List Bullet")

        document.add_heading("4. Danh sách giao dịch", level=2)

        table = document.add_table(rows=1, cols=6)

        table.style = "Table Grid"

        headers = ["ID", "Loại", "Ngày", "Danh mục", "Nội dung", "Số tiền"]

        for index, header in enumerate(headers):

            table.rows[0].cells[index].text = header

        if transactions:

            sorted_transactions = sorted(transactions, key=lambda t: self.model.parse_date(t.get("date", "")) or datetime.min)

            for t in sorted_transactions:

                row = table.add_row().cells

                row[0].text = str(t.get("id", ""))

                row[1].text = str(t.get("type", ""))

                row[2].text = str(t.get("date", ""))

                row[3].text = str(t.get("category", "Khác"))

                row[4].text = str(t.get("description", ""))

                row[5].text = f"{int(self.model.to_float(t.get('amount', '0'))):,} đ"

        else:

            row = table.add_row().cells

            row[0].text = "Không có giao dịch nào trong bộ lọc hiện tại"

        for paragraph in document.paragraphs:

            for run in paragraph.runs:

                run.font.name = "Arial"

                run.font.size = Pt(11)

        document.save(file_path)

        return file_path

    def export_report_to_pdf(self, file_path, transactions=None):

        from reportlab.lib.pagesizes import A4

        from reportlab.lib import colors

        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        from reportlab.lib.units import cm

        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

        from reportlab.pdfbase import pdfmetrics

        from reportlab.pdfbase.ttfonts import TTFont

        font_name = "Helvetica"

        bold_font_name = "Helvetica-Bold"

        possible_fonts = [

            r"C:\Windows\Fonts\arial.ttf",

            r"C:\Windows\Fonts\tahoma.ttf",

            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"

        ]

        possible_bold_fonts = [

            r"C:\Windows\Fonts\arialbd.ttf",

            r"C:\Windows\Fonts\tahomabd.ttf",

            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

        ]

        for font_path in possible_fonts:

            if os.path.exists(font_path):

                pdfmetrics.registerFont(TTFont("UnicodeFont", font_path))

                font_name = "UnicodeFont"

                break

        for font_path in possible_bold_fonts:

            if os.path.exists(font_path):

                pdfmetrics.registerFont(TTFont("UnicodeFontBold", font_path))

                bold_font_name = "UnicodeFontBold"

                break

        if transactions is None:

            transactions = self.model.get_all_transactions()

        total_income = sum(self.model.to_float(t["amount"]) for t in transactions if t.get("type") == "Thu nhập")

        total_expense = sum(self.model.to_float(t["amount"]) for t in transactions if t.get("type") == "Chi tiêu")

        balance = total_income - total_expense

        category_percentages = self.model.get_expense_percentage_by_category(transactions)

        monthly_expenses = self.model.get_expenses_by_month(transactions)

        document = SimpleDocTemplate(

            file_path,

            pagesize=A4,

            rightMargin=1.5 * cm,

            leftMargin=1.5 * cm,

            topMargin=1.5 * cm,

            bottomMargin=1.5 * cm

        )

        styles = getSampleStyleSheet()

        styles.add(ParagraphStyle(name="VNTitle", fontName=bold_font_name, fontSize=18, alignment=1, spaceAfter=14))

        styles.add(ParagraphStyle(name="VNHeading", fontName=bold_font_name, fontSize=13, spaceBefore=10, spaceAfter=8))

        styles.add(ParagraphStyle(name="VNNormal", fontName=font_name, fontSize=10, leading=14))

        story = []

        story.append(Paragraph("BÁO CÁO THU CHI CÁ NHÂN", styles["VNTitle"]))

        story.append(Paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", styles["VNNormal"]))

        story.append(Spacer(1, 8))

        story.append(Paragraph("1. Tổng quan tài chính", styles["VNHeading"]))

        story.append(Paragraph(f"Tổng thu nhập: {int(total_income):,} đ", styles["VNNormal"]))

        story.append(Paragraph(f"Tổng chi tiêu: {int(total_expense):,} đ", styles["VNNormal"]))

        story.append(Paragraph(f"Số dư còn lại: {int(balance):,} đ", styles["VNNormal"]))

        story.append(Paragraph("2. Chi tiêu theo danh mục", styles["VNHeading"]))

        if category_percentages:

            for category, (amount, percent) in sorted(category_percentages.items(), key=lambda x: x[1][0], reverse=True):

                story.append(Paragraph(f"- {category}: {int(amount):,} đ ({percent}%)", styles["VNNormal"]))

        else:

            story.append(Paragraph("- Chưa có giao dịch chi tiêu", styles["VNNormal"]))

        story.append(Paragraph("3. Chi tiêu theo tháng", styles["VNHeading"]))

        if monthly_expenses:

            for month, amount in sorted(monthly_expenses.items(), key=lambda x: datetime.strptime(x[0], "%m/%Y")):

                story.append(Paragraph(f"- {month}: {int(amount):,} đ", styles["VNNormal"]))

        else:

            story.append(Paragraph("- Chưa có giao dịch chi tiêu", styles["VNNormal"]))

        story.append(Paragraph("4. Danh sách giao dịch", styles["VNHeading"]))

        table_data = [["ID", "Loại", "Ngày", "Danh mục", "Nội dung", "Số tiền"]]

        if transactions:

            sorted_transactions = sorted(transactions, key=lambda t: self.model.parse_date(t.get("date", "")) or datetime.min)

            for t in sorted_transactions:

                table_data.append([

                    str(t.get("id", "")),

                    str(t.get("type", "")),

                    str(t.get("date", "")),

                    str(t.get("category", "Khác")),

                    str(t.get("description", "")),

                    f"{int(self.model.to_float(t.get('amount', '0'))):,} đ"

                ])

        else:

            table_data.append(["Không có giao dịch", "", "", "", "", ""])

        report_table = Table(table_data, repeatRows=1, colWidths=[1.1*cm, 2.0*cm, 2.2*cm, 2.5*cm, 5.5*cm, 3.0*cm])

        report_table.setStyle(TableStyle([

            ("FONTNAME", (0, 0), (-1, 0), bold_font_name),

            ("FONTNAME", (0, 1), (-1, -1), font_name),

            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),

            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),

            ("VALIGN", (0, 0), (-1, -1), "TOP"),

            ("ALIGN", (5, 1), (5, -1), "RIGHT"),

            ("FONTSIZE", (0, 0), (-1, -1), 8),

            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),

            ("TOPPADDING", (0, 0), (-1, -1), 6),

        ]))

        story.append(report_table)

        document.build(story)

        return file_path

    def export_report_to_excel(self, transactions=None):

        if transactions is None:

            transactions = self.model.get_all_transactions()

        total_income = sum(self.model.to_float(t["amount"]) for t in transactions if t.get("type") == "Thu nhập")

        total_expense = sum(self.model.to_float(t["amount"]) for t in transactions if t.get("type") == "Chi tiêu")

        balance = total_income - total_expense

        category_percentages = self.model.get_expense_percentage_by_category(transactions)

        monthly_expenses = self.model.get_expenses_by_month(transactions)

        os.makedirs("reports", exist_ok=True)

        file_name = f"bao_cao_thu_chi_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        file_path = os.path.join("reports", file_name)

        summary_data = [

            {"Metric": "Tổng thu nhập", "Value": int(total_income)},

            {"Metric": "Tổng chi tiêu", "Value": int(total_expense)},

            {"Metric": "Số dư còn lại", "Value": int(balance)}

        ]

        category_data = [

            {"Danh mục": category, "Số tiền": int(amount), "Tỷ lệ (%)": percent}

            for category, (amount, percent) in sorted(category_percentages.items(), key=lambda x: x[1][0], reverse=True)

        ]

        monthly_data = [

            {"Tháng": month, "Số tiền": int(amount)}

            for month, amount in sorted(monthly_expenses.items(), key=lambda x: datetime.strptime(x[0], "%m/%Y"))

        ]

        transactions_data = [

            {

                "ID": t.get("id", ""),

                "Loại": t.get("type", ""),

                "Ngày": t.get("date", ""),

                "Danh mục": t.get("category", "Khác"),

                "Nội dung": t.get("description", ""),

                "Số tiền": int(self.model.to_float(t.get("amount", "0")))

            }

            for t in sorted(transactions, key=lambda t: self.model.parse_date(t.get("date", "")) or datetime.min)

        ]

        with pd.ExcelWriter(file_path, engine="openpyxl") as writer:

            pd.DataFrame(summary_data).to_excel(writer, sheet_name="Tổng quan", index=False)

            pd.DataFrame(category_data).to_excel(writer, sheet_name="Theo danh mục", index=False)

            pd.DataFrame(monthly_data).to_excel(writer, sheet_name="Theo tháng", index=False)

            pd.DataFrame(transactions_data).to_excel(writer, sheet_name="Giao dịch", index=False)

        return file_path
